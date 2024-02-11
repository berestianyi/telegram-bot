from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from function.utils import name_split, name_cut, month_names, \
    car_data, car_quantity, today_date, date_to_dict, is_man_for_fop, is_man_for_tov

user_data = {"atp_choice": "roland_h", "name": "Деркач Кирило Сергійович",
             "expire_date": "10.10.2024",
             "vehicle_data": "1) BO0424EB, MERCEDES-BENZ SPRINTER 519 CDI, 2011; "
                             "2) BO6360EB,  RENAULT TRAFIC, 2017; "
                             "3) BO5153EO,  RENAULT MASTER, 2011; "
                             "4) BO3001CA, БОГДАН А-09212, 2008;",
             "tov_name": "Деркача Кирила Сергійовича",
             "company": "«ВЕСТ ТРАНС БУС»",
             "code": "1245347400", "address": "04202, Україна, 04202, місто Київ, вулиця Кондратюка Юрія 5, будинок "
                                              "5, квартира 563",
             "email": "anton_reva140@ukr.net", "phone": "+38(067)-966-15-15"}


quan = car_quantity(user_data["vehicle_data"])

# print(car_data(user_data["vehicle_data"], quan))


def test_roland_doc(user_info):

    if len(user_info["code"]) == 10:
        is_tov = False
        if user_info["atp_choice"] == "roland_d":
            doc = Document('roland_d_fop.docx')
        else:
            doc = Document('roland_h_fop.docx')
        if is_man_for_fop(user_info["code"]):
            sho = "який"
        else:
            sho = "яка"
    else:
        is_tov = True
        if user_info["atp_choice"] == "roland_d":
            doc = Document('roland_d_tov.docx')
        else:
            doc = Document('roland_h_tov.docx')
        if is_man_for_tov(user_info['name']):
            sho = "який"
        else:
            sho = "яка"

    first_table = doc.tables[0]
    second_table = doc.tables[1]

    half_name = name_cut(user_info['name'], "half_name")
    small_name = name_cut(user_info['name'], "small_name")

    quantity = car_quantity(user_data["vehicle_data"])
    cars_data = car_data(user_data["vehicle_data"], quan)

    for i in range(1, quantity + 1):
        cells = first_table.add_row().cells
        cells[0].paragraphs[0].add_run(str(i)).bold = True
        cells[1].paragraphs[0].add_run(cars_data["number" + str(i)]).bold = True
        cells[2].paragraphs[0].add_run(cars_data["name" + str(i)]).bold = True
        cells[3].paragraphs[0].add_run(cars_data["year" + str(i)]).bold = True

    for row in first_table.rows:
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    today = date_to_dict(today_date())
    expire = date_to_dict(user_info['expire_date'])

    doc_num = today["dd"] + "/" + today["mm"]

    dictionary = {"dd": today["dd"],
                  "mm": month_names[today["mm"]],
                  "yyyy": today["yyyy"]}

    for i in dictionary:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i, dictionary[i])

    doc.paragraphs[0].add_run("ДОГОВІР № " + doc_num).bold = True

    if is_tov:
        doc.paragraphs[6].add_run("ТОВАРИСТВО З ОБМЕЖЕНОЮ ВІДПОВІДАЛЬНІСТЮ " + user_info["company"].upper() +
                                  " в особі Директора " + user_info["name"].upper() + ", ").bold = True
        doc.paragraphs[6].add_run(sho + " діє на підставі Статуту (надалі – Замовник) з другої сторони, "
                                  "в подальшому разом іменуються «Сторони», а кожна окремо – «Сторона», "
                                  "керуючись чинним законодавством України та нормативно-правовими документами, "
                                  "які регламентують діяльність підприємства автомобільного транспорту всіх форм "
                                  "власності, уклали даний Договір про наступне:")
    else:
        doc.paragraphs[6].add_run("ФІЗИЧНА ОСОБА-ПІДПРИЄМЕЦЬ " + user_info["name"].upper() + ", ").bold = True
        doc.paragraphs[6].add_run(sho + " діє на підставі виписки із ЄДРПОУ (надалі – Замовник) з другої сторони, в "
                                  "подальшому разом іменуються «Сторони», а кожна окремо – «Сторона», керуючись "
                                  "чинним законодавством України та нормативно-правовими документами, які "
                                  "регламентують діяльність підприємства автомобільного транспорту всіх "
                                  "форм власності, уклали даний "
                                  "Договір про наступне:").bold = False

    doc.paragraphs[35].add_run("4.1. ").bold = True
    doc.paragraphs[35].add_run("Даний Договір вступає в силу з ").bold = False
    doc.paragraphs[35].add_run(today["dd"] + " " + month_names[today["mm"]]
                               + " " + today["yyyy"] + " року і діє до" + " " +
                               expire["dd"] + " " +
                               month_names[expire["mm"]] + " " + expire["yyyy"]
                               + " року.").bold = True

    if is_tov:
        name_table = second_table.rows[1].cells[1].paragraphs[1]
        name_table.add_run(user_info['company'].upper()).bold = True
    else:
        name_table = second_table.rows[1].cells[1].paragraphs[1]
        name_table.add_run(user_info['name'].upper()).bold = True

    index_address_table = second_table.rows[2].cells[1].paragraphs[0]
    index_address_table.text = user_info['address']

    code_edpoy_table = second_table.rows[2].cells[1].paragraphs[1]
    code_edpoy_table.add_run("Код ЄДРПОУ " + user_info["code"])

    if user_info['phone'] is not None:
        phone_table = second_table.rows[2].cells[1].paragraphs[2]
        phone_table.add_run("Тел. " + user_info['phone'])

    if user_info['email'] is not None:
        email_table = second_table.rows[2].cells[1].paragraphs[3]
        email_table.add_run("Е-mail: " + user_info["email"])

    name2_table = second_table.rows[3].cells[1].paragraphs[3]
    name2_table.add_run("____________________ " + half_name.lower().title()).bold = True

    for row in second_table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.name = 'Arial'

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    doc_name = 'Д АТП ФОП ' + small_name + '.docx'
    doc.save(doc_name)

    return doc_name


# tov_roland_doc(user_data)

def test_royal_tov(user_info):
    is_tov = True

    if len(user_info["code"]) == 10:
        is_tov = False

    if is_tov:
        doc = Document('royal_tov.docx')
        if is_man_for_tov(user_info['name']):
            sho = "який"
        else:
            sho = "яка"
    else:
        doc = Document('royal_fop.docx')
        if is_man_for_fop(user_info["code"]):
            sho = "який"
        else:
            sho = "яка"

    top_table = doc.tables[0]
    first_table = doc.tables[1]
    second_table = doc.tables[2]

    first_name = name_split(user_info['name'], "first_name")
    surname = name_split(user_info['name'], "surname")
    small_name = name_cut(user_info['name'], "small_name")

    today = date_to_dict(today_date())
    expire = date_to_dict(user_info['expire_date'])

    quantity = car_quantity(user_data["vehicle_data"])
    cars_data = car_data(user_data["vehicle_data"], quan)

    for i in range(1, quantity + 1):
        cells = first_table.add_row().cells
        cells[0].paragraphs[0].add_run(str(i)).bold = True
        cells[1].paragraphs[0].add_run(cars_data["number" + str(i)]).bold = True
        cells[2].paragraphs[0].add_run(cars_data["name" + str(i)]).bold = True
        cells[3].paragraphs[0].add_run(cars_data["year" + str(i)]).bold = True

    for row in first_table.rows:
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc_num = today["dd"] + "/" + today["mm"]

    dictionary = {"dd": today["dd"],
                  "mm": month_names[today["mm"]],
                  "yyyy": today["yyyy"]}

    for i in dictionary:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i, dictionary[i])

    name_table = top_table.rows[0].cells[0].paragraphs[1]
    name_table.add_run("ДОГОВІР № " + doc_num).bold = True

    if is_tov:
        doc.paragraphs[3].add_run("ТОВАРИСТВО З ОБМЕЖЕНОЮ ВІДПОВІДАЛЬНІСТЮ " + user_info["company"].upper()).bold = True
        doc.paragraphs[3].add_run(" в особі Директора ").bold = False
        doc.paragraphs[3].add_run(user_info["name"].upper() + ", ").bold = True
        doc.paragraphs[3].add_run(sho + " діє на підставі виписки із ЄДРПОУ (надалі іменується – Замовник), з "
                                  "однієї сторони та ").bold = False
    else:
        doc.paragraphs[3].add_run("ФІЗИЧНА ОСОБА-ПІДПРИЄМЕЦЬ " + user_info["name"].upper() + ", ").bold = True
        doc.paragraphs[3].add_run(sho + " діє на підставі виписки із ЄДРПОУ (надалі іменується – Замовник), з "
                                  "однієї сторони та ").bold = False

    doc.paragraphs[40].add_run("4.1. ").bold = True
    doc.paragraphs[40].add_run("Даний Договір вступає в силу з ").bold = False
    doc.paragraphs[40].add_run(today["dd"] + " " + month_names[today["mm"]]
                               + " " + today["yyyy"] + " року і діє до" + " " +
                               expire["dd"] + " " +
                               month_names[expire["mm"]] + " " + expire["yyyy"]
                               + " року.").bold = True

    if is_tov:
        name_table = second_table.rows[0].cells[0].paragraphs[3]
        run = name_table.add_run(user_info['company'].upper())
        run.font.bold = True
        run.font.underline = True
    else:
        name_table = second_table.rows[0].cells[0].paragraphs[3]
        run = name_table.add_run(user_info['name'].upper())
        run.font.bold = True
        run.font.underline = True

    index_address_table = second_table.rows[1].cells[0].paragraphs[1]
    index_address_table.text = user_info['address']

    code_edpoy_table = second_table.rows[1].cells[0].paragraphs[0]
    code_edpoy_table.add_run("Код ЄДРПОУ: " + user_info["code"])

    if user_info['phone'] is not None:
        phone_table = second_table.rows[1].cells[0].paragraphs[2]
        phone_table.add_run("Тел. " + user_info['phone'])

    if user_info['email'] is not None:
        email_table = second_table.rows[1].cells[0].paragraphs[3]
        email_table.add_run("Е-mail: " + user_info["email"])

    if is_tov:
        name2_table = second_table.rows[2].cells[0].paragraphs[0]
        name2_table.add_run(user_info["company"]).bold = True

    name2_table = second_table.rows[2].cells[0].paragraphs[3]
    name2_table.add_run("____________________ " + first_name.lower().title() + ' ' + surname.upper()).bold = True

    for row in second_table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.name = 'Times New Roman'

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    doc_name = 'Д АТП ФОП ' + small_name + '.docx'
    doc.save(doc_name)

    return doc_name


test_roland_doc(user_data)
